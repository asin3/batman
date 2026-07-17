"""
============================================================
Batman DOX Import Engine (BDIE)

Document Parser

Purpose:
Inspects a document and reports its structure.

This is NOT a curriculum parser.
It simply tells Batman what exists inside the document.

============================================================
"""

from docx.document import Document
from src.platform.dox.schema_registry import normalize
from src.platform.dox.profiles.profile_detector import (
    ProfileDetector,
)

############################################################
# DOCUMENT PARSER
############################################################

class DocumentParser:

    def __init__(self, document: Document):

        self.document = document

        self.profile_detector = ProfileDetector()

    ########################################################
    # GROUPS
    ########################################################

    def get_groups(self):

        groups = {}

        context = self.get_context()

        for index, table in enumerate(self.document.tables):

            headers = [
                normalize(cell.text)
                for cell in table.rows[0].cells
            ]

            profile = self.profile_detector.detect(headers)

            if profile is None:
                continue

            if profile.name != "GroupedCurriculumProfile":
                continue

            subject = context[index]["subject"]

            groups[subject] = profile.get_groups(
                table,
                headers,
            )

        return groups
    
    ########################################################
    # CHAPTERS
    ########################################################

    def get_chapters(self):

        chapters = {}

        context = self.get_context()

        for index, table in enumerate(self.document.tables):

            headers = [
                normalize(cell.text)
                for cell in table.rows[0].cells
            ]

            profile = self.profile_detector.detect(headers)

            if profile is None:

                continue

            if profile.name not in (
                "GroupedCurriculumProfile",
                "FlatCurriculumProfile",
            ):

                continue

            subject = context[index]["subject"]

            chapters[subject] = profile.get_chapters(
                table,
                headers,
            )

        return chapters
    ########################################################
    # TOPICS
    ########################################################

    def get_topics(self):

        topics = {}

        context = self.get_context()

        for index, table in enumerate(self.document.tables):

            headers = [
                normalize(cell.text)
                for cell in table.rows[0].cells
            ]
            profile = self.profile_detector.detect(headers)

            if profile is None:
                continue

            subject = context[index]["subject"]

            topics[subject] = profile.get_topics(
                table,
                headers,
            )

        return topics

########################################################
# DOCUMENT CONTEXT
########################################################

    def get_context(self):

        context = []

        default_subjects = [
            "Mathematics",
            "Physics",
            "Chemistry",
            "Biology",
        ]

        for index, table in enumerate(self.document.tables):

            item = {

                "table_index": index,

                "subject": (
                    default_subjects[index]
                    if index < len(default_subjects)
                    else f"Subject {index + 1}"
                ),

                "rows": len(table.rows),

                "columns": len(table.columns)

            }

            context.append(item)

        return context

    ########################################################
    # PARSE TABLE
    ########################################################

    def parse_table(self, table):

        headers = [
            normalize(cell.text)
            for cell in table.rows[0].cells
        ]

        profile = self.profile_detector.detect(headers)

        if profile is None:

            return None

        return profile.parse(
            table,
            headers,
        )

    ########################################################
    # SUMMARY
    ########################################################

    def summary(self):

        print("\n==============================")
        print("Document Summary")
        print("==============================")

        print(f"Paragraphs : {len(self.document.paragraphs)}")
        print(f"Tables     : {len(self.document.tables)}")

        heading_count = 0

        for paragraph in self.document.paragraphs:

            if paragraph.style.name.startswith("Heading"):
                heading_count += 1

        print(f"Headings   : {heading_count}")

        print("==============================\n")

        self.table_summary()

        self.analyze_structure()

        self.normalize_schema()


    ########################################################
    # TABLE SUMMARY
    ########################################################

    def table_summary(self):

        print("\n==============================")
        print("Table Analysis")
        print("==============================")

        for index, table in enumerate(self.document.tables, start=1):

            rows = len(table.rows)
            cols = len(table.columns)

            print(f"\nTable {index}")
            print(f"Rows    : {rows}")
            print(f"Columns : {cols}")

            if rows > 0:

                headers = []

                for cell in table.rows[0].cells:
                    headers.append(cell.text.strip())

                print("Headers :")

                for header in headers:
                    print(f"  - {header}")

        print("\n==============================")

########################################################
# ANALYZE STRUCTURE
########################################################

    def analyze_structure(self):

        print("\n==============================")
        print("Structure Analysis")
        print("==============================")

        for index, table in enumerate(self.document.tables, start=1):

            if len(table.rows) == 0:
                continue

            headers = [
                cell.text.strip().lower()
                for cell in table.rows[0].cells
            ]

            structure = "Unknown"
            confidence = "Low"

            # ----------------------------------------
            # Curriculum Table
            # ----------------------------------------

            if (
                any("chapter" in h for h in headers)
                and any("unit" in h for h in headers)
            ):
                structure = "Curriculum Table"
                confidence = "High"

            elif any("chapter" in h for h in headers):
                structure = "Chapter Table"
                confidence = "Medium"

            print(f"\nTable {index}")
            print(f"Detected   : {structure}")
            print(f"Confidence : {confidence}")

        print("\n==============================")

########################################################
# NORMALIZE SCHEMA
########################################################

    def normalize_schema(self):

        print("\n==============================")
        print("Schema Normalization")
        print("==============================")

        for index, table in enumerate(self.document.tables, start=1):

            if len(table.rows) == 0:
                continue

            normalized = []

            for cell in table.rows[0].cells:

                normalized.append(
                    normalize(cell.text)
                )

            print(f"\nTable {index}")

            for column in normalized:
                print(f"  • {column}")

        print("\n==============================")

############################################################
# TEST
############################################################

if __name__ == "__main__":

    print("Batman Document Parser Ready")